from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import json
import os
import sys
import docx
import shutil
import io
import csv
from typing import List, Dict, Any

# Ensure backend directory is in path
sys.path.append(os.path.dirname(__file__))
from agents import (
    JobIntelligenceAgent,
    CandidateProfileAgent,
    CareerIntelligenceAgent,
    SkillIntelligenceAgent,
    BehaviorIntelligenceAgent,
    CulturalFitAgent,
    HoneypotAgent,
    ScoringAgent,
    RankingAgent,
    ReasoningAgent
)

app = FastAPI(title="Redrob AI Candidate Intelligence Engine API")

# Enable CORS for the React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Try to load the default job description text on startup
DEFAULT_JD_TEXT = ""
DEFAULT_JD_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Dataset", "job_description.docx")
if os.path.exists(DEFAULT_JD_PATH):
    try:
        doc = docx.Document(DEFAULT_JD_PATH)
        paragraphs = [p.text for p in doc.paragraphs]
        DEFAULT_JD_TEXT = "\n".join(paragraphs)
    except Exception as e:
        print(f"Error loading default JD text: {e}")

# Load precalculated data if available (e.g. for Hugging Face Spaces instant demo)
PRECALCULATED_RESULTS_PATH = os.path.join(os.path.dirname(__file__), "precalculated_results.json")
PRECALCULATED_STATS_PATH = os.path.join(os.path.dirname(__file__), "precalculated_stats.json")

def load_precalculated_data():
    initial_ranked_results = []
    initial_stats = None
    if os.path.exists(PRECALCULATED_RESULTS_PATH) and os.path.exists(PRECALCULATED_STATS_PATH):
        try:
            with open(PRECALCULATED_RESULTS_PATH, "r", encoding="utf-8") as f:
                initial_ranked_results = json.load(f)
            with open(PRECALCULATED_STATS_PATH, "r", encoding="utf-8") as f:
                initial_stats = json.load(f)
        except Exception as e:
            print(f"Error loading precalculated data: {e}")
    return initial_ranked_results, initial_stats

initial_ranked_results, initial_stats = load_precalculated_data()

# In-memory storage for current state
CURRENT_STATE = {
    "jd_text": DEFAULT_JD_TEXT,
    "jd_spec": JobIntelligenceAgent().analyze_jd(DEFAULT_JD_TEXT) if DEFAULT_JD_TEXT else {},
    "candidates": [], # Raw candidate objects
    "ranked_results": initial_ranked_results, # Scored and ranked candidates
    "stats": initial_stats, # Precalculated stats
    "weights": {
        "w_tech": 0.40,
        "w_career": 0.25,
        "w_culture": 0.20,
        "w_education": 0.15
    },
    "progress": {"status": "idle", "percent": 0, "message": ""}
}

# Ensure temp directory exists
TEMP_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "temp")
os.makedirs(TEMP_DIR, exist_ok=True)

class WeightsModel(BaseModel):
    w_tech: float
    w_career: float
    w_culture: float
    w_education: float

@app.get("/api/config")
def get_config():
    return CURRENT_STATE["weights"]

@app.post("/api/config")
def update_config(weights: WeightsModel):
    CURRENT_STATE["weights"] = {
        "w_tech": weights.w_tech,
        "w_career": weights.w_career,
        "w_culture": weights.w_culture,
        "w_education": weights.w_education
    }
    # Re-score and re-rank if we have candidates loaded
    if CURRENT_STATE["candidates"]:
        run_ranking_in_memory()
    return {"message": "Weights updated successfully", "weights": CURRENT_STATE["weights"]}

@app.post("/api/upload-jd")
async def upload_jd(file: UploadFile = File(...)):
    filename = file.filename
    content = await file.read()
    
    jd_text = ""
    if filename.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs]
            jd_text = "\n".join(paragraphs)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse docx: {str(e)}")
    elif filename.endswith(".txt") or filename.endswith(".md"):
        jd_text = content.decode("utf-8")
    else:
        raise HTTPException(status_code=400, detail="Only .docx, .txt, or .md files are supported")
        
    CURRENT_STATE["jd_text"] = jd_text
    
    # Initialize job agent and extract
    job_agent = JobIntelligenceAgent()
    CURRENT_STATE["jd_spec"] = job_agent.analyze_jd(jd_text)
    
    # Reset precalculated state since user is starting a live demo run
    CURRENT_STATE["ranked_results"] = []
    CURRENT_STATE["stats"] = None
    
    return {
        "message": "Job description uploaded successfully",
        "length": len(jd_text),
        "spec": CURRENT_STATE["jd_spec"]
    }

@app.get("/api/jd")
def get_jd():
    return {
        "jd_text": CURRENT_STATE["jd_text"],
        "jd_spec": CURRENT_STATE["jd_spec"]
    }

@app.get("/api/stats")
def get_stats_endpoint():
    if CURRENT_STATE["stats"]:
        return CURRENT_STATE["stats"]
    if CURRENT_STATE["ranked_results"]:
        return get_stats()
    return None

@app.post("/api/reset")
def reset_to_precalculated():
    initial_ranked_results, initial_stats = load_precalculated_data()
    CURRENT_STATE["jd_text"] = DEFAULT_JD_TEXT
    CURRENT_STATE["jd_spec"] = JobIntelligenceAgent().analyze_jd(DEFAULT_JD_TEXT) if DEFAULT_JD_TEXT else {}
    CURRENT_STATE["candidates"] = []
    CURRENT_STATE["ranked_results"] = initial_ranked_results
    CURRENT_STATE["stats"] = initial_stats
    CURRENT_STATE["weights"] = {
        "w_tech": 0.40,
        "w_career": 0.25,
        "w_culture": 0.20,
        "w_education": 0.15
    }
    CURRENT_STATE["progress"] = {"status": "idle", "percent": 0, "message": ""}
    return {"message": "Reset to 100K precalculated report successfully"}

@app.post("/api/upload-candidates")
async def upload_candidates(file: UploadFile = File(...)):
    filename = file.filename
    
    # Reset candidates and clear precalculated stats
    CURRENT_STATE["candidates"] = []
    CURRENT_STATE["ranked_results"] = []
    CURRENT_STATE["stats"] = None
    
    candidates_list = []
    try:
        # Stream line-by-line from the file object to optimize memory usage
        for line_bytes in file.file:
            line = line_bytes.decode("utf-8").strip()
            if not line:
                continue
            try:
                # Handle both JSON array elements and JSONL lines
                if line.startswith("[") or line.startswith(","):
                    clean_line = line.lstrip("[").rstrip("]").rstrip(",")
                    if not clean_line:
                        continue
                    candidates_list.append(json.loads(clean_line))
                else:
                    candidates_list.append(json.loads(line))
            except Exception:
                # If a single line fails to parse, it could be a full JSON array file
                # We raise an error so the fallback block below handles it
                raise ValueError("Line parsing failed")
    except Exception:
        # Fallback: If line-by-line streaming failed or it's a formatted JSON array file,
        # try loading the entire file as a single JSON object.
        try:
            file.file.seek(0)
            content = file.file.read()
            candidates_list = json.loads(content.decode("utf-8"))
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse candidates data: {str(e)}")
                
    CURRENT_STATE["candidates"] = candidates_list
    return {
        "message": f"Successfully uploaded {len(candidates_list)} candidates",
        "count": len(candidates_list)
    }

@app.get("/api/progress")
def get_progress():
    return CURRENT_STATE["progress"]

@app.post("/api/rank")
def trigger_ranking():
    if not CURRENT_STATE["candidates"]:
        raise HTTPException(status_code=400, detail="No candidates loaded. Upload candidates first.")
        
    CURRENT_STATE["progress"] = {"status": "running", "percent": 10, "message": "Initializing agents..."}
    try:
        run_ranking_in_memory()
        CURRENT_STATE["progress"] = {"status": "done", "percent": 100, "message": "Ranking completed"}
        return {
            "message": "Ranking complete",
            "top_100": CURRENT_STATE["ranked_results"][:100],
            "stats": get_stats()
        }
    except Exception as e:
        CURRENT_STATE["progress"] = {"status": "error", "percent": 0, "message": f"Error: {str(e)}"}
        raise HTTPException(status_code=500, detail=f"Ranking failed: {str(e)}")

def run_ranking_in_memory():
    candidates = CURRENT_STATE["candidates"]
    weights = CURRENT_STATE["weights"]
    
    # Initialize agents
    job_agent = JobIntelligenceAgent()
    jd_spec = CURRENT_STATE["jd_spec"] if CURRENT_STATE["jd_spec"] else job_agent.analyze_jd()
    
    profile_agent = CandidateProfileAgent()
    career_agent = CareerIntelligenceAgent(jd_spec)
    skill_agent = SkillIntelligenceAgent(jd_spec)
    behavior_agent = BehaviorIntelligenceAgent()
    culture_agent = CulturalFitAgent()
    honeypot_agent = HoneypotAgent()
    scoring_agent = ScoringAgent()
    
    # Update scoring weights
    scoring_agent.weights = weights
    
    ranking_agent = RankingAgent()
    reasoning_agent = ReasoningAgent()
    
    candidates_scored = []
    total = len(candidates)
    update_interval = max(1, total // 10)
    
    for idx, c in enumerate(candidates):
        profile = profile_agent.normalize(c)
        risk_data = honeypot_agent.detect_risk(profile)
        career_data = career_agent.analyze_career(profile)
        skill_data = skill_agent.evaluate_skills(profile)
        behavior_data = behavior_agent.analyze_behavior(profile)
        culture_score = culture_agent.evaluate_culture(profile, career_data)
        
        scores = scoring_agent.compute_score(
            tech_data=skill_data,
            career_data=career_data,
            behavior_data=behavior_data,
            culture_score=culture_score,
            edu_data={"education": profile["education"]},
            risk_data=risk_data
        )
        
        candidates_scored.append({
            "candidate_id": profile["candidate_id"],
            "profile": profile,
            "scores": scores,
            "career_data": career_data,
            "behavior_data": behavior_data,
            "risk_data": risk_data
        })
        
        if (idx + 1) % update_interval == 0 or (idx + 1) == total:
            percent = 10 + int(75 * (idx + 1) / total)
            CURRENT_STATE["progress"] = {
                "status": "running",
                "percent": percent,
                "message": f"Scoring candidates ({idx + 1}/{total})..."
            }
        
    top_100 = ranking_agent.rank_candidates(candidates_scored)
    
    # Generate explanations
    CURRENT_STATE["progress"] = {
        "status": "running",
        "percent": 90,
        "message": "Generating explanations for top 100..."
    }
    for item in top_100:
        reasoning = reasoning_agent.generate_reasoning(
            profile=item["profile"],
            scores=item["scores"],
            career_analysis=item["career_data"],
            behavior_analysis=item["behavior_data"],
            risk_analysis=item["risk_data"],
            rank=item["rank"]
        )
        item["reasoning"] = reasoning
        
    CURRENT_STATE["ranked_results"] = top_100

def get_stats():
    ranked = CURRENT_STATE["ranked_results"]
    candidates = CURRENT_STATE["candidates"]
    
    total = len(candidates)
    
    # Count honeypots in whole dataset
    # We can approximate or run it for stats
    honeypots = 0
    job_agent = JobIntelligenceAgent()
    jd_spec = CURRENT_STATE["jd_spec"] if CURRENT_STATE["jd_spec"] else job_agent.analyze_jd()
    profile_agent = CandidateProfileAgent()
    honeypot_agent = HoneypotAgent()
    
    for c in candidates[:5000]:  # Cap at 5000 for quick stats if dataset is huge
        p = profile_agent.normalize(c)
        if honeypot_agent.detect_risk(p)["is_honeypot"]:
            honeypots += 1
            
    avg_score = sum(item["scores"]["final_score"] for item in ranked) / len(ranked) if ranked else 0.0
    
    locations = {}
    skills_dist = {}
    
    for item in ranked[:100]:
        loc = item["profile"]["location"]
        locations[loc] = locations.get(loc, 0) + 1
        
        for s in item["profile"]["skills"][:5]:
            sname = s["name"]
            skills_dist[sname] = skills_dist.get(sname, 0) + 1
            
    return {
        "total_candidates": total,
        "sample_honeypot_ratio": round((honeypots / min(total, 5000)) * 100, 2) if total > 0 else 0,
        "top_100_avg_score": round(avg_score, 4),
        "locations": sorted([{"name": k, "count": v} for k, v in locations.items()], key=lambda x: -x["count"])[:5],
        "skills": sorted([{"name": k, "count": v} for k, v in skills_dist.items()], key=lambda x: -x["count"])[:10]
    }

@app.get("/api/leaderboard")
def get_leaderboard():
    if not CURRENT_STATE["ranked_results"]:
        return []
    
    # Return lightweight representation of top 100
    leaderboard = []
    for item in CURRENT_STATE["ranked_results"]:
        leaderboard.append({
            "rank": item["rank"],
            "candidate_id": item["candidate_id"],
            "name": item["profile"]["name"],
            "title": item["profile"]["current_title"],
            "company": item["profile"]["current_company"],
            "years_of_experience": item["profile"]["years_of_experience"],
            "score": item["scores"]["final_score"],
            "reasoning": item["reasoning"]
        })
    return leaderboard

@app.get("/api/candidates/{candidate_id}")
def get_candidate(candidate_id: str):
    found = None
    for item in CURRENT_STATE["ranked_results"]:
        if item["candidate_id"] == candidate_id:
            found = item
            break
            
    if not found:
        # Check raw candidates if not in top 100
        for c in CURRENT_STATE["candidates"]:
            if c.get("candidate_id") == candidate_id:
                # Score them on the fly
                job_agent = JobIntelligenceAgent()
                jd_spec = CURRENT_STATE["jd_spec"] if CURRENT_STATE["jd_spec"] else job_agent.analyze_jd()
                profile_agent = CandidateProfileAgent()
                career_agent = CareerIntelligenceAgent(jd_spec)
                skill_agent = SkillIntelligenceAgent(jd_spec)
                behavior_agent = BehaviorIntelligenceAgent()
                culture_agent = CulturalFitAgent()
                honeypot_agent = HoneypotAgent()
                scoring_agent = ScoringAgent()
                scoring_agent.weights = CURRENT_STATE["weights"]
                
                profile = profile_agent.normalize(c)
                risk_data = honeypot_agent.detect_risk(profile)
                career_data = career_agent.analyze_career(profile)
                skill_data = skill_agent.evaluate_skills(profile)
                behavior_data = behavior_agent.analyze_behavior(profile)
                culture_score = culture_agent.evaluate_culture(profile, career_data)
                
                scores = scoring_agent.compute_score(
                    tech_data=skill_data,
                    career_data=career_data,
                    behavior_data=behavior_data,
                    culture_score=culture_score,
                    edu_data={"education": profile["education"]},
                    risk_data=risk_data
                )
                
                reasoning_agent = ReasoningAgent()
                reasoning = reasoning_agent.generate_reasoning(
                    profile=profile,
                    scores=scores,
                    career_analysis=career_data,
                    behavior_analysis=behavior_data,
                    risk_analysis=risk_data,
                    rank=101
                )
                
                found = {
                    "candidate_id": candidate_id,
                    "profile": profile,
                    "scores": scores,
                    "career_data": career_data,
                    "behavior_data": behavior_data,
                    "risk_data": risk_data,
                    "rank": "Outside Top 100",
                    "reasoning": reasoning
                }
                break
                
    if not found:
        raise HTTPException(status_code=404, detail="Candidate not found")
        
    return found

@app.get("/api/download-csv")
def download_csv():
    if not CURRENT_STATE["ranked_results"]:
        raise HTTPException(status_code=400, detail="No ranking results available. Run ranking first.")
        
    import pandas as pd
    excel_path = os.path.join(TEMP_DIR, "export_submission.xlsx")
    
    data = []
    for item in CURRENT_STATE["ranked_results"]:
        data.append({
            "candidate_id": item["candidate_id"],
            "rank": item["rank"],
            "score": round(item["scores"]["final_score"], 4),
            "reasoning": item["reasoning"]
        })
    df = pd.DataFrame(data)
    df.to_excel(excel_path, index=False)
            
    return FileResponse(
        path=excel_path,
        filename="TeamBrahmastra_IndiaRunsRedRobAI.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# Serve React static assets
frontend_dist_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend", "dist")
root_dist_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "dist")
if os.path.exists(frontend_dist_path):
    app.mount("/", StaticFiles(directory=frontend_dist_path, html=True), name="static")
elif os.path.exists(root_dist_path):
    app.mount("/", StaticFiles(directory=root_dist_path, html=True), name="static")

